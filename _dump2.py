from docx import Document
d = Document('docs/EI_Manuscript_FINAL.docx')
for i in [7, 51, 64, 92, 93, 116, 117]:
    print(f"--- PARA {i} ---")
    print(repr(d.paragraphs[i].text))
    print()
print("--- §4.5 context ---")
for i in range(108, 115):
    print(i, repr(d.paragraphs[i].text[:600]))
