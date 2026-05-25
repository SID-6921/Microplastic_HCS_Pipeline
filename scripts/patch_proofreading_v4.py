"""Round-2 proofreading patcher.

Fixes:
  A) Para 72 in docs/EI_Manuscript_FINAL.docx -- corrupted superscript.
     "p < 10\u207b\u00b9\u2077\u00b0" (degree sign) -> "p < 10\u207b\u00b9\u2077\u2070" (superscript zero).
  B) Table 8 (Table 7 in MS numbering), rows 1/4/7 column 2:
     "o (100 nm)" -> "nano (100 nm)" (PS, PE, PET).
"""
from __future__ import annotations
import sys
from pathlib import Path
from docx import Document

DOCX = Path("docs/EI_Manuscript_FINAL.docx")

SUPERSCRIPT_OLD = "p\u202f<\u202f10\u207b\u00b9\u2077\u00b0"
SUPERSCRIPT_NEW = "p\u202f<\u202f10\u207b\u00b9\u2077\u2070"

CELL_OLD = "o (100 nm)"
CELL_NEW = "nano (100 nm)"


def replace_in_paragraph(p, old: str, new: str) -> bool:
    if old not in p.text:
        return False
    full = p.text.replace(old, new)
    # collapse runs, preserve first-run formatting
    for r in p.runs[1:]:
        r.text = ""
    if p.runs:
        p.runs[0].text = full
    else:
        p.add_run(full)
    return True


def main() -> int:
    if not DOCX.exists():
        print(f"[FATAL] {DOCX} not found")
        return 2
    doc = Document(str(DOCX))
    hits = 0

    # A) superscript
    for i, p in enumerate(doc.paragraphs):
        if SUPERSCRIPT_OLD in p.text:
            if replace_in_paragraph(p, SUPERSCRIPT_OLD, SUPERSCRIPT_NEW):
                print(f"[OK]   superscript fix in paragraph {i}")
                hits += 1
    if hits == 0:
        print("[MISS] superscript pattern not found")

    # B) table cells
    cell_hits = 0
    for ti, t in enumerate(doc.tables):
        for ri, row in enumerate(t.rows):
            for ci, cell in enumerate(row.cells):
                if CELL_OLD in cell.text and "nano" not in cell.text:
                    # replace inside cell paragraphs
                    for p in cell.paragraphs:
                        if replace_in_paragraph(p, CELL_OLD, CELL_NEW):
                            print(f"[OK]   table {ti} row {ri} col {ci}: {CELL_OLD!r} -> {CELL_NEW!r}")
                            cell_hits += 1
    if cell_hits == 0:
        print("[MISS] no table cells updated")

    doc.save(str(DOCX))
    print(f"\nSaved {DOCX}. paragraph_hits={hits}, cell_hits={cell_hits}")
    return 0 if (hits + cell_hits) > 0 else 1


if __name__ == "__main__":
    sys.exit(main())
