"""Round-3 proofreading patcher (must-fix #2,3,4 + should-fix #5,6,7,8).

Idempotent in-place DOCX edits.
"""
from __future__ import annotations
import sys
from pathlib import Path
from docx import Document

DOCX = Path("docs/EI_Manuscript_FINAL.docx")


def replace_in_paragraph(p, old: str, new: str) -> bool:
    if old not in p.text:
        return False
    full = p.text.replace(old, new)
    for r in p.runs[1:]:
        r.text = ""
    if p.runs:
        p.runs[0].text = full
    else:
        p.add_run(full)
    return True


# === #3 Title (finding-led, Option B) ============================================
TITLE_OLD = (
    "Feature-Based Models Outperform Deep Learning in Simulation-Derived "
    "High-Content Screening Data: A Benchmark Revealing Dominant-Feature "
    "Dependency in Cell-Death Phenotype Classification"
)
TITLE_NEW = (
    "Logistic Regression Achieves Superior Cross-Validated Stability and "
    "Calibration Over Deep Learning in Simulation-Derived High-Content "
    "Screening Data: A Dominant-Feature Dependency Benchmark for Cell-Death "
    "Phenotype Classification"
)

# === #5 Abstract: tone down LR stability claim ===================================
ABSTRACT_OLD = "LR nonetheless retained superior cross-fold stability, calibration"
ABSTRACT_NEW = (
    "LR nonetheless retained marginally superior cross-fold stability "
    "(CV AUC 0.981\u202f\u00b1\u202f0.004 vs 0.977\u202f\u00b1\u202f0.005), "
    "calibration"
)

# === #6 \u00a73.4 RF ablation: use 0.872 consistently (Table 10) =====================
ABLATION_OLD = "reduces AUC from 0.972 to 0.870 (\u0394\u202f=\u2009\u22120.103)"
ABLATION_NEW = "reduces AUC from 0.972 to 0.872 (\u0394\u202f=\u2009\u22120.100)"

# === #2 Cite [17] in \u00a74.5 (Caicedo 2022 nucleus segmentation) ===================
S45_OLD = (
    "CellProfiler [5] and related tools extract overlapping morphological "
    "descriptors from real HCS images."
)
S45_NEW = (
    "CellProfiler [5] and related tools extract overlapping morphological "
    "descriptors from real HCS images; upstream nucleus segmentation - a "
    "prerequisite for any HCS profiling pipeline - has been substantially "
    "advanced by community benchmarks such as the 2018 Data Science Bowl, "
    "which established modern reference methods for this step [17]."
)

# === #4 Table 9 caption: add cross-regime warning ===============================
TBL9_OLD = (
    "Table 9. Original permutation tests (RF as reference; retained for "
    "completeness)."
)
TBL9_NEW = (
    "Table 9. Original permutation tests (RF as reference; retained for "
    "completeness). Note: AUC values in this table are from the single "
    "held-out test split (n\u202f=\u202f200) and are not directly comparable "
    "to the cross-validation macro-AUC values in Table 9b."
)

# === #7 Table 11 note: drop [27,28] =============================================
TBL11_OLD = "and from cited literature [8,9,19,27,28]"
TBL11_NEW = "and from cited literature [8,9,19]"

# === #8 Table 2 note: explain RF CI upper bound 1.000 ===========================
TBL2_OLD = (
    "Macro-AUC 95% CI: analytic Hanley-McNeil approximation "
    "(n_pos\u202f=\u202f50, n_neg\u202f=\u202f150)."
)
TBL2_NEW = (
    "Macro-AUC 95% CI: analytic Hanley-McNeil approximation "
    "(n_pos\u202f=\u202f50, n_neg\u202f=\u202f150). The RF upper CI bound of "
    "1.000 reflects truncation of the analytic Hanley-McNeil approximation "
    "at the boundary; bootstrap CIs would provide more reliable interval "
    "estimates at this sample size."
)


REPLACEMENTS = [
    ("#3 title (finding-led)",      TITLE_OLD,    TITLE_NEW),
    ("#5 abstract LR stability",    ABSTRACT_OLD, ABSTRACT_NEW),
    ("#6 \u00a73.4 RF ablation 0.870->0.872", ABLATION_OLD, ABLATION_NEW),
    ("#2 \u00a74.5 cite [17]",          S45_OLD,      S45_NEW),
    ("#4 Table 9 cross-regime note", TBL9_OLD,     TBL9_NEW),
    ("#7 Table 11 note refs",       TBL11_OLD,    TBL11_NEW),
    ("#8 Table 2 CI=1.000 note",    TBL2_OLD,     TBL2_NEW),
]


def main() -> int:
    if not DOCX.exists():
        print(f"[FATAL] {DOCX} not found")
        return 2
    doc = Document(str(DOCX))

    print("[PATCH report]")
    total_hits = 0
    for label, old, new in REPLACEMENTS:
        hits = 0
        for p in doc.paragraphs:
            if replace_in_paragraph(p, old, new):
                hits += 1
        # also check table cells just in case
        for t in doc.tables:
            for row in t.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        if replace_in_paragraph(p, old, new):
                            hits += 1
        tag = "[OK]  " if hits > 0 else "[MISS]"
        print(f"  {tag} {label}: {hits} replacement(s)")
        total_hits += hits

    doc.save(str(DOCX))
    print(f"[DONE] saved {DOCX.resolve()}; total replacements={total_hits}")
    return 0 if total_hits > 0 else 1


if __name__ == "__main__":
    sys.exit(main())
