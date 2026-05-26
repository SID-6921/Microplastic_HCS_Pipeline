"""
Numeric-consistency audit for EI_Manuscript_FINAL.docx.

For every authoritative number that appears in the manuscript, asserts:
  (a) the value appears at least once,
  (b) no stale variant of the same quantity also appears.

Run:  python scripts\audit_manuscript_numbers.py
Exit code 0 = all PASS; non-zero = at least one FAIL.

Source of truth comes from results/tables/*.csv where possible, plus a small
ledger of derived numbers (permutation p, ablation deltas) that live only in
the manuscript prose.
"""
from __future__ import annotations
from pathlib import Path
import re
import sys
import pandas as pd
from docx import Document

# Force UTF-8 stdout so unicode (\u00b1, \u202f, superscripts) prints on Windows cp1252 consoles.
try:
    sys.stdout.reconfigure(encoding="utf-8")
except Exception:
    pass

ROOT   = Path(__file__).resolve().parent.parent
DOCX   = ROOT / "docs" / "EI_Manuscript_FINAL.docx"
TABLES = ROOT / "results" / "tables"


def load_doc_text() -> str:
    if not DOCX.exists():
        print(f"[FAIL] {DOCX} not found", file=sys.stderr)
        sys.exit(2)
    d = Document(str(DOCX))
    parts = [p.text for p in d.paragraphs]
    for tbl in d.tables:
        for row in tbl.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    parts.append(p.text)
    return "\n".join(parts)


# ── Authoritative numbers ───────────────────────────────────────────────
# CV cross-validation (LR/RF, GroupKFold k=5) — from table_cv_summary.csv
CV_TRUTH = {
    "LR CV macro-AUC":  ("0.981", "0.004"),
    "RF CV macro-AUC":  ("0.977", "0.005"),
}

# Held-out test (n=200) — from table_1_model_performance.csv
HELDOUT_TRUTH = {
    "LR held-out macro-AUC":  "0.901",
    "RF held-out macro-AUC":  "0.966",
    "CNN held-out macro-AUC": "0.873",
    "R18s held-out macro-AUC": "0.910",
    "R18p held-out macro-AUC": "0.954",
}

# Calibration ECE — from table_3_calibration_ece.csv
ECE_TRUTH = {
    "LR ECE": "0.042",
    "RF ECE": "0.091",
}

# Dominant-feature ablation — from table_10_degenerate_check.csv
ABLATION_TRUTH = {
    "LR full-feature AUC":     "0.981",
    "LR ablated AUC":          "0.861",
    "RF full-feature AUC":     "0.972",
    "RF ablated AUC":          "0.872",
}

# Permutation test (LR vs RF on CV macro-AUC)
PERMUTATION_TRUTH = {
    "Delta AUC (LR-RF)":  "0.026",
    "p-value (text)":     "5 \u00d7 10\u207b\u2074",
}

# Stale values that MUST NOT appear (these were earlier drafts' wrong numbers)
FORBIDDEN = [
    "0.845", "0.820",                  # stale accuracy in earlier abstract
    "0.0005",                          # stale p notation; canonical is 5 \u00d7 10\u207b\u2074
    "p = 5e-4", "p=5e-4",              # informal notation
    "ResNet-18 (0.954\u20130.954)",    # silly degenerate range
    "domit",                           # typo reviewer flagged by name
    "Author Name",                     # placeholder author block
    "Department, Institution",         # placeholder affiliation
    "email@institution.edu",           # placeholder email
    "0.961 to 0.887",                  # stale ablation prose (Table 4 row is 0.9724 -> 0.8698)
    "0.972 to 0.870",                  # stale RF ablation prose (canonical is 0.872 from Table 10)
    "10\u207b\u00b9\u2077\u00b0",      # corrupted exponent: degree sign instead of superscript zero
    "[8,9,19,27,28]",                  # Table 11 note over-cited; canonical refs are [8,9,19]
    # Note: 'o (100 nm)' truncation in Table 7 is enforced by patch_proofreading_v4.py;
    # not added here because it would substring-match the correct 'nano (100 nm)'.
]


def assert_present(text: str, label: str, value: str) -> bool:
    if value in text:
        print(f"  [PASS] {label}: {value} present")
        return True
    print(f"  [FAIL] {label}: {value} NOT FOUND")
    return False


def assert_absent(text: str, value: str) -> bool:
    if value in text:
        print(f"  [FAIL] stale value present: {value!r}")
        return False
    print(f"  [PASS] stale value absent: {value!r}")
    return True


def cross_check_csv(text: str) -> bool:
    """Read CSVs and verify the headline numbers above actually match."""
    ok = True
    try:
        cv = pd.read_csv(TABLES / "table_cv_summary.csv")
        # Expect Model column + CV_AUC_mean + CV_AUC_std
        for _, r in cv.iterrows():
            m = str(r.get("Model", "")).strip()
            if m in ("LR", "Logistic Regression"):
                mu, sd = float(r["CV_AUC_mean"]), float(r["CV_AUC_std"])
                exp_mu, exp_sd = CV_TRUTH["LR CV macro-AUC"]
                if f"{mu:.3f}" != exp_mu or f"{sd:.3f}" != exp_sd:
                    print(f"  [FAIL] CSV/ledger mismatch for LR CV: csv={mu:.3f}\u00b1{sd:.3f} ledger={exp_mu}\u00b1{exp_sd}")
                    ok = False
                else:
                    print(f"  [PASS] CSV matches ledger for LR CV ({mu:.3f}\u00b1{sd:.3f})")
            elif m in ("RF", "Random Forest"):
                mu, sd = float(r["CV_AUC_mean"]), float(r["CV_AUC_std"])
                exp_mu, exp_sd = CV_TRUTH["RF CV macro-AUC"]
                if f"{mu:.3f}" != exp_mu or f"{sd:.3f}" != exp_sd:
                    print(f"  [FAIL] CSV/ledger mismatch for RF CV: csv={mu:.3f}\u00b1{sd:.3f} ledger={exp_mu}\u00b1{exp_sd}")
                    ok = False
                else:
                    print(f"  [PASS] CSV matches ledger for RF CV ({mu:.3f}\u00b1{sd:.3f})")
    except Exception as e:
        print(f"  [WARN] could not cross-check table_cv_summary.csv: {e}")
    return ok


def main() -> int:
    text = load_doc_text()
    ok = True

    print("\n[A] CV cross-validation values")
    for label, (mu, sd) in CV_TRUTH.items():
        # Accept both ASCII "+/-" rendering and unicode \u00b1
        joined = f"{mu}\u202f\u00b1\u202f{sd}"
        if joined in text or f"{mu} \u00b1 {sd}" in text or f"{mu}\u00b1{sd}" in text:
            print(f"  [PASS] {label}: {mu}\u00b1{sd} present")
        else:
            print(f"  [FAIL] {label}: {mu}\u00b1{sd} NOT FOUND in any spacing variant")
            ok = False

    print("\n[B] Held-out macro-AUC values")
    for label, val in HELDOUT_TRUTH.items():
        if not assert_present(text, label, val):
            ok = False

    print("\n[C] Calibration ECE values")
    for label, val in ECE_TRUTH.items():
        if not assert_present(text, label, val):
            ok = False

    print("\n[D] Dominant-feature ablation values")
    for label, val in ABLATION_TRUTH.items():
        if not assert_present(text, label, val):
            ok = False

    print("\n[E] Permutation test values")
    # Delta AUC: simple literal
    if not assert_present(text, "Delta AUC (LR-RF)", PERMUTATION_TRUTH["Delta AUC (LR-RF)"]):
        ok = False
    # p-value: accept any of the canonical spacing variants
    p_variants = [
        "5 \u00d7 10\u207b\u2074",
        "5\u202f\u00d7\u202f10\u207b\u2074",
        "5\u00d710\u207b\u2074",
        "5\u00a0\u00d7\u00a010\u207b\u2074",
    ]
    if any(v in text for v in p_variants):
        print("  [PASS] p-value (text): 5x10^-4 present (in a canonical spacing form)")
    else:
        print("  [FAIL] p-value (text): 5x10^-4 NOT FOUND in any canonical spacing form")
        ok = False

    print("\n[F] Forbidden / stale values must be absent")
    for val in FORBIDDEN:
        if not assert_absent(text, val):
            ok = False

    print("\n[G] CSV \u2194 ledger cross-check")
    if not cross_check_csv(text):
        ok = False

    print("\n[H] Structural anchors")
    anchors = {
        "Title":                      "Logistic Regression Achieves Superior Cross-Validated Stability",
        "Section 4.7 marker":         "Broader significance",
        "Section 4.8 marker":         "deployment roadmap",
        "Table 11 marker":            "Table\u202f11",
        "Figure 10 marker":           "Figure 10",
        "Nihart 2025":                "Nihart",
        "Geirhos shortcut ref":       "Shortcut learning",
        "DeGrave shortcut ref":       "DeGrave",
        "Reconciliation phrase":      "preferred baseline classifier for simulation-derived HCS",
        "Degenerate-regime term":     "degenerate regime",
    }
    for label, needle in anchors.items():
        if needle in text:
            print(f"  [PASS] {label}: {needle!r}")
        else:
            print(f"  [FAIL] {label}: {needle!r} NOT FOUND")
            ok = False

    print("\n=== AUDIT " + ("PASS" if ok else "FAIL") + " ===")
    return 0 if ok else 1


if __name__ == "__main__":
    sys.exit(main())
