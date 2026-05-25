"""
In-place DOCX proofreading patcher for EI_Manuscript_FINAL.docx.

Applies the six reviewer-flagged fixes plus two sentence-flow splits:
  (1) 'Without domit feature'  -> 'Without dominant feature'           (Table 10, x2)
  (2) Author/affiliation placeholders -> Siddhardha Nanda + TBD affil
  (3) §3.4 RF ablation '0.961 to 0.887 (\u0394 = -0.074)'
                       -> '0.972 to 0.870 (\u0394 = -0.103)'           (matches Table 4)
  (4) §4.7 'vs.\u202f  DL'    -> 'vs.\u202f DL'                        (double space)
  (5) §3.1 long sentence split (Taken together \u2026)
  (6) §4.1 long sentence split (RF=0.966, R18p=0.954, LR=0.901)

Idempotent: safe to re-run.
"""
from __future__ import annotations
from pathlib import Path
import sys
from docx import Document

ROOT = Path(__file__).resolve().parent.parent
DOCX = ROOT / "docs" / "EI_Manuscript_FINAL.docx"

# ------------------------------------------------------------------ helpers
def replace_in_paragraph(p, old: str, new: str) -> bool:
    """Replace old->new inside a paragraph. If old spans runs, collapses runs."""
    full = "".join(r.text for r in p.runs)
    if old not in full:
        return False
    new_full = full.replace(old, new)
    # collapse: keep first run, clear text on others
    if p.runs:
        p.runs[0].text = new_full
        for r in p.runs[1:]:
            r.text = ""
    else:
        p.add_run(new_full)
    return True


def replace_in_cell(cell, old: str, new: str) -> bool:
    """Cells contain paragraphs; replace inside each paragraph."""
    hit = False
    for p in cell.paragraphs:
        if replace_in_paragraph(p, old, new):
            hit = True
    return hit


# ------------------------------------------------------------------ patches

# 1) "domit" -> "dominant" inside every table cell + every paragraph.
DOMIT_FIX = ("Without domit feature", "Without dominant feature")

# 2) Author / affiliation placeholders.
AUTHOR_OLD = "Author Name\u00b9, Author Name\u00b2, Corresponding Author\u00b3"
AUTHOR_NEW = "Siddhardha Nanda\u00b9"
AFFIL_OLD = (
    "\u00b9 Department, Institution, City, Country  "
    "\u00b2 Department, Institution, City, Country  "
    "\u00b3 Corresponding author: email@institution.edu"
)
AFFIL_NEW = (
    "\u00b9 Independent Researcher  "
    "\u2709 Corresponding author: [affiliation and email to be added at submission]"
)

# 3) Section 3.4 RF ablation numeric reconciliation with Table 4 row 0/1.
# Exact bytes in DOCX use narrow no-break (\u202f) around '=' and thin space (\u2009) before minus.
ABLATION_OLD = "reduces AUC from 0.961 to 0.887 (\u0394\u202f=\u2009\u22120.074)"
ABLATION_NEW = "reduces AUC from 0.972 to 0.870 (\u0394\u202f=\u2009\u22120.103)"

# 4) §4.7 double-space after 'vs.\u202f'.
DOUBLESPACE_OLD = "vs.\u202f DL ECE"
DOUBLESPACE_NEW = "vs.\u202fDL ECE"

# 5) §3.1 long sentence split. Exact DOCX uses U+202F around ±, =, ×, units, etc.
S31_OLD = (
    "Although RF achieved a higher held-out macro-AUC than LR on this single "
    "test split, LR demonstrated superior stability across GroupKFold "
    "cross-validation (0.981\u202f\u00b1\u202f0.004 vs 0.977\u202f\u00b1\u202f0.005), "
    "substantially better calibration (ECE 0.042 vs 0.091, Table\u202f3), "
    "lower training-time cost (0.02\u202fs vs 0.48\u202fs, Table\u202f6), and a "
    "significant CV-level advantage in the permutation test "
    "(\u0394AUC\u202f=\u202f0.026, p\u202f=\u202f5\u202f\u00d7\u202f10\u207b\u2074, Table\u202f9b)."
)
S31_NEW = (
    "Although RF achieved a higher held-out macro-AUC than LR on this single "
    "test split, LR demonstrated superior stability across GroupKFold "
    "cross-validation (0.981\u202f\u00b1\u202f0.004 vs 0.977\u202f\u00b1\u202f0.005) "
    "and substantially better calibration (ECE 0.042 vs 0.091, Table\u202f3). "
    "LR also incurred lower training-time cost (0.02\u202fs vs 0.48\u202fs, "
    "Table\u202f6) and held a significant CV-level advantage in the permutation "
    "test (\u0394AUC\u202f=\u202f0.026, p\u202f=\u202f5\u202f\u00d7\u202f10\u207b\u2074, Table\u202f9b)."
)

# 6) §4.1 long sentence split.
S41_OLD = (
    "On the single 200-sample held-out split the ranking partially inverts "
    "(RF\u202f=\u202f0.966, R18p\u202f=\u202f0.954, LR\u202f=\u202f0.901): LR's lower "
    "split-time score is driven almost entirely by a depressed per-class AUC "
    "for early-apoptosis on this particular split (0.665 vs.\u202f0.939 for RF; "
    "Table\u202f2), an effect consistent with the higher between-split variance "
    "expected from a single 200-sample held-out evaluation and resolved by the "
    "GroupKFold estimate."
)
S41_NEW = (
    "On the single 200-sample held-out split the ranking partially inverts "
    "(RF\u202f=\u202f0.966, R18p\u202f=\u202f0.954, LR\u202f=\u202f0.901). LR's "
    "lower split-time score is driven almost entirely by a depressed per-class "
    "AUC for early-apoptosis on this split (0.665 vs.\u202f0.939 for RF; "
    "Table\u202f2). This pattern is consistent with the higher between-split "
    "variance expected from a single 200-sample evaluation and is resolved by "
    "the GroupKFold estimate."
)

PARA_FIXES = [
    ("author block",        AUTHOR_OLD,      AUTHOR_NEW),
    ("affil block",         AFFIL_OLD,       AFFIL_NEW),
    ("\u00a73.4 ablation",  ABLATION_OLD,    ABLATION_NEW),
    ("\u00a74.7 dblspace",  DOUBLESPACE_OLD, DOUBLESPACE_NEW),
    ("\u00a73.1 split",     S31_OLD,         S31_NEW),
    ("\u00a74.1 split",     S41_OLD,         S41_NEW),
]


def main() -> int:
    if not DOCX.exists():
        print(f"[FAIL] {DOCX} not found", file=sys.stderr)
        return 2

    d = Document(str(DOCX))
    log = []

    # Table-cell fix: "domit" -> "dominant"
    domit_hits = 0
    for t in d.tables:
        for r in t.rows:
            for c in r.cells:
                if replace_in_cell(c, *DOMIT_FIX):
                    domit_hits += 1
    log.append(("\u00a73.2 Table 10 'domit' \u2192 'dominant'", domit_hits))

    # Paragraph fixes
    for label, old, new in PARA_FIXES:
        hits = 0
        for p in d.paragraphs:
            if replace_in_paragraph(p, old, new):
                hits += 1
        log.append((label, hits))

    # Save
    d.save(str(DOCX))

    print("[PATCH report]")
    fail = False
    for label, hits in log:
        status = "OK" if hits > 0 else "MISS"
        if hits == 0:
            fail = True
        print(f"  [{status}] {label}: {hits} replacement(s)")

    if fail:
        # 'MISS' is informational; only fail hard if NOTHING patched
        if all(h == 0 for _, h in log):
            return 1
    print(f"[DONE] saved {DOCX}")
    return 0


if __name__ == "__main__":
    sys.exit(main())
