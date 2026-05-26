from docx import Document
import re
d = Document('docs/EI_Manuscript_FINAL.docx')
print("=== PARAGRAPHS 90-100 ===")
for i in range(90, 101):
    print(f"--- {i} ---")
    print(repr(d.paragraphs[i].text))
print()
print("=== ALL TABLE CAPTIONS / Notes anywhere ===")
for i,p in enumerate(d.paragraphs):
    if re.match(r'^(Table\s|Note:)', p.text):
        print(i, repr(p.text[:300]))
print()
print("=== Para 51 (Table 2 note) ===")
print(repr(d.paragraphs[51].text))
print()
print("=== [27] [28] usage ===")
for i,p in enumerate(d.paragraphs):
    if '[27' in p.text or '[28' in p.text:
        print(i, repr(p.text[:400]))
print()
print("=== Table 9 (not 9b) area ===")
for i,p in enumerate(d.paragraphs):
    if 'Table 9' in p.text and '9b' not in p.text and i not in [92,93,95,96]:
        print(i, repr(p.text[:300]))
