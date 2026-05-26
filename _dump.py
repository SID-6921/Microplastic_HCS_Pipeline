from docx import Document
d = Document('docs/EI_Manuscript_FINAL.docx')
for i,p in enumerate(d.paragraphs):
    if i < 10: print(i, repr(p.text[:200]))
print("---")
for i,p in enumerate(d.paragraphs):
    t=p.text
    if 'Caicedo' in t or '[17]' in t or 'cross-fold stability' in t or 'Abstract' in t or t.startswith('4.5') or t.startswith('3.4 ') or t.startswith('4.5 '):
        print(i, repr(t[:400]))
print("---NOTES NEAR TABLES---")
for i,p in enumerate(d.paragraphs):
    t=p.text
    if ('Note' in t or 'note:' in t.lower() or 'Table 2' in t or 'Table\u202f2' in t or 'Table 9' in t or 'Table\u202f9' in t or 'Table 11' in t or 'Table\u202f11' in t) and len(t)<800:
        print(i, repr(t[:400]))
